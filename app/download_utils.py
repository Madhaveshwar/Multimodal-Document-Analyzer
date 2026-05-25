"""Generate downloadable TXT, PDF, and DOCX files from summary text."""

from __future__ import annotations

import io
from datetime import datetime
from typing import Iterable


def to_txt(summary: str) -> bytes:
    """Return summary as UTF-8 encoded plain text bytes."""
    return summary.encode("utf-8")


def to_pdf(summary: str, title: str = "Document Summary") -> bytes:
    """Return summary as a PDF byte stream using ReportLab."""
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "Title", parent=styles["Heading1"], fontSize=16, spaceAfter=12
    )
    body_style = ParagraphStyle(
        "Body", parent=styles["Normal"], fontSize=11, leading=16
    )

    # Escape HTML-special chars for ReportLab Paragraph
    safe_summary = (
        summary.replace("&", "&amp;")
               .replace("<", "&lt;")
               .replace(">", "&gt;")
               .replace("\n", "<br/>")
    )

    story = [
        Paragraph(title, title_style),
        Spacer(1, 0.3 * cm),
        Paragraph(safe_summary, body_style),
    ]
    doc.build(story)
    return buf.getvalue()


def to_docx(summary: str, title: str = "Document Summary") -> bytes:
    """Return summary as a DOCX byte stream using python-docx."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(title, level=1)
    para = doc.add_paragraph(summary)
    para.runs[0].font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _safe_timestamp(value: str | None) -> str:
    if not value:
        return "unknown"
    try:
        return datetime.fromisoformat(value.replace("Z", "+00:00")).astimezone().strftime("%Y-%m-%d %H:%M:%S %Z")
    except Exception:
        return value


def _format_chat_messages(messages: Iterable[dict]) -> str:
    lines: list[str] = []
    for index, message in enumerate(messages, start=1):
        role = str(message.get("role", "message")).title()
        text = str(message.get("text", "")).strip()
        created_at = _safe_timestamp(message.get("created_at"))
        lines.append(f"{index}. {role} [{created_at}]")
        if text:
            lines.append(text)
        sources = message.get("sources") or []
        if sources:
            lines.append("Sources:")
            for source_index, source in enumerate(sources, start=1):
                filename = source.get("filename", "unknown")
                page_number = source.get("page_number")
                page_label = f"p.{page_number}" if isinstance(page_number, int) and page_number > 0 else "p.n/a"
                citation = source.get("citation", source_index)
                chunk_source = source.get("chunk_source", "unknown")
                lines.append(f"  [{citation}] {filename} {page_label} - {chunk_source}")
                source_text = str(source.get("text", "")).strip()
                if source_text:
                    lines.append(f"    {source_text}")
        lines.append("")
    return "\n".join(lines).strip()


def chat_history_to_txt(messages: Iterable[dict], title: str = "Conversation Export") -> bytes:
    """Return a chat transcript as UTF-8 bytes."""
    transcript = _format_chat_messages(messages)
    header = f"{title}\n{'=' * len(title)}\n\n"
    return f"{header}{transcript}\n".encode("utf-8")


def chat_history_to_pdf(messages: Iterable[dict], title: str = "Conversation Export") -> bytes:
    """Return a chat transcript as a PDF byte stream."""
    transcript = _format_chat_messages(messages)
    return to_pdf(transcript, title=title)


def chat_history_to_docx(messages: Iterable[dict], title: str = "Conversation Export") -> bytes:
    """Return a chat transcript as a DOCX byte stream."""
    transcript = _format_chat_messages(messages)
    return to_docx(transcript, title=title)
